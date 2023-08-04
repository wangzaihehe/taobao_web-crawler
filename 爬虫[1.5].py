import pyautogui
import sys
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from bs4 import BeautifulSoup
import time
import cv2
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.support.ui import  WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import re

from docx import Document
import docx
import os
import requests as req
from PIL import Image
import io
from io import BytesIO
from docx.shared import Inches
import imageio
import magic
import numpy as np

import configparser

import random
def text_progress_bar(total, current, length=50):
    percent = current / total
    arrow = '=' * int(length * percent)
    spaces = ' ' * (length - len(arrow))
    sys.stdout.write(f"\r[{arrow}{spaces}] {int(percent * 100)}%")
    sys.stdout.flush()

def longest_common_substring(str1, str2, excluded_str):
    m, n = len(str1), len(str2)
    # 创建一个二维数组用于保存中间结果
    dp = [[0] * (n + 1) for _ in range(m + 1)]
    max_len = 0  # 用于保存最长公共子串的长度
    end_index = 0  # 用于保存最长公共子串的结束索引

    for i in range(1, m + 1):
        for j in range(1, n + 1):
            if str1[i - 1] == str2[j - 1]:
                dp[i][j] = dp[i - 1][j - 1] + 1
                if dp[i][j] > max_len and str1[i - dp[i][j]:i] != excluded_str:
                    max_len = dp[i][j]
                    end_index = i

    return str1[end_index - max_len:end_index]

def find_lc(str1:str,str2:str):
    """
    找到str1和str2之间的最长子字符串，返回其长度和该字符串
    """
    m = [ [ 0 for _ in range(len(str2)+1) ] for _ in range(len(str1)+1) ]  #长度
    d = [ [ '' for _ in range(len(str2)+1) ] for _ in range(len(str1)+1) ]  #子字符串

    for p1 in range(len(str1)):
        for p2 in range(len(str2)):
            if str1[p1]==str2[p2]:
                m[p1+1][p2+1]=m[p1][p2]+1
                d[p1+1][p2+1]=d[p1][p2]+str1[p1]
            elif m[p1+1][p2]>m[p1][p2+1]:
                m[p1+1][p2+1]=m[p1+1][p2]
                d[p1+1][p2+1]=d[p1+1][p2]
            else:
                m[p1+1][p2+1]=m[p1][p2+1]
                d[p1+1][p2+1]=d[p1][p2+1]
    
    return (m[len(str1)][len(str2)])



def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = docx.text.run.Run(
        docx.oxml.shared.OxmlElement('w:r'), paragraph)
    new_run.text = text

    # Set the run's style to the builtin hyperlink style, defining it if necessary
    # new_run.style = get_or_create_hyperlink_style(part.document)
    # Alternatively, set the run's formatting explicitly
    new_run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
    new_run.font.underline = True

    # Join all the xml elements together
    hyperlink.append(new_run._element)  
    paragraph._p.append(hyperlink)
    return hyperlink


# chrome_img = cv2.imread(r"C:\Users\WIN10\Desktop\chrome1.png") 
# pyautogui.hotkey('win','m')
# location_chrome = pyautogui.locateOnScreen(chrome_img,confidence = 0.7)
# print(location_chrome)
# pyautogui.doubleClick(location_chrome)
# searchbar_img = cv2.imread(r"C:\Users\WIN10\Desktop\searchbar.png") 
# location_searchbar = pyautogui.locateOnScreen(searchbar_img ,confidence = 0.7)
# print(location_searchbar)
# pyautogui.click(location_searchbar)
# pyautogui.write('https://login.taobao.com/')

# docx_name = pyautogui.prompt(text='请输入位于桌面的文档名「无需后缀」',title='购物车爬虫消息框',default='')
# id = pyautogui.prompt(text='请输入淘宝用户名',title='购物车爬虫消息框',default='')
# password = pyautogui.password(text='输入密码',title='购物车爬虫消息框',default='',mask='*')
# pyautogui.alert(text='在登陆后如遇二重验证,请在15秒内完成手机确认',title='购物车爬虫消息框',button='OK')

current_file_path = os.path.abspath(__file__)
current_folder = os.path.dirname(current_file_path)
general_cfg_path = os.path.join(current_folder,'General.cfg')
config = configparser.ConfigParser()
config.read(general_cfg_path,encoding='utf-8')
id =config.get('General', 'id')
password = config.get('General', 'password')
docx_name = config.get('General','docx_name')
option = Options()
option.add_experimental_option('excludeSwitches', ['enable-automation'])
option.add_argument('--disable-blink-features=AutomationControlled')
bro=webdriver.Chrome(options = option)
bro.maximize_window()
bro.implicitly_wait(5)
wait = WebDriverWait(bro, 10)
bro.get('https://login.taobao.com/')
# login_img = cv2.imread(r"C:\Users\WIN10\Desktop\username.png") 
# location_login = pyautogui.locateOnScreen(login_img)
# print(location_login)

# pyautogui.typewrite('Hello world')
# driver = webdriver.Chrome()

bro.find_element(By.ID,'fm-login-id').click()
bro.find_element(By.ID,'fm-login-id').clear()
bro.find_element(By.ID,'fm-login-id').send_keys(id)

bro.find_element(By.ID,'fm-login-password').clear()
bro.find_element(By.ID,'fm-login-password').send_keys(password)
bro.find_element(By.ID,'fm-login-password').send_keys(Keys.ENTER)
time.sleep(4)
# bro.find_element(By.CLASS_NAME,'fm-button.fm-submit.password-login').click()
#滑块验证
path_verification = os.path.abspath(os.path.dirname(__file__))
# print(path_verification)
# sys.stdout.flush()
slider_img = cv2.imread(path_verification+"\dragslider.png") 
try:

    location_slider = pyautogui.locateOnScreen(slider_img ,confidence = 0.8)
    
    point_slider = pyautogui.center(location_slider)
    time.sleep(4)
    x, y = point_slider

    pyautogui.moveTo(x,y)
    pyautogui.dragTo(x+300, y, 0.5, button='left') 
except:
    # sys.stdout.write("没有滑动条")
    time.sleep(1)

time.sleep(3)
handle=bro.current_window_handle
bro.find_element(By.ID,'J_MiniCart').click()
bro.implicitly_wait(5)
wait = WebDriverWait(bro, 10)
#滚动
scroll_pause_time = 5 
last_height = bro.execute_script('return document.body.scrollHeight')
while True:
    bro.find_element('tag name', 'body').send_keys(Keys.END)
    time.sleep(scroll_pause_time)
    new_height = bro.execute_script('return document.body.scrollHeight')
    if new_height == last_height:
        break
    last_height = new_height
bro.implicitly_wait(5)
#全部界面加载完成
# try:
#     # 打开网页
#     url = "https://detail.tmall.com/item.htm?id=617990296152&skuId=4640067052817&spm=a1z0d.6639537/tb.1997196601.98.16f47484cFInQX"  # 将此处替换为您想要访问的网页地址
#     bro.get(url)

#     # 获取整个网页的HTML内容
#     html_content = bro.page_source

#     # 将HTML内容转换为字符串
#     html_string = str(html_content)

#     # 要搜索的目标字符串
#     target_string = "skuItem"
#     target_string2 = "tb-txt"

#     # 检查目标字符串是否在HTML中存在
#     if target_string in html_string:
#         # 在HTML中找到目标字符串，模拟鼠标点击该元素

#         print("天猫")
#         sys.stdout.flush()
#     elif target_string2 in html_string:
#         print("淘宝")
#         sys.stdout.flush()
#     else:
#         print("目标字符串在HTML中未找到。")
# finally:
#     # 关闭浏览器
#     bro.quit()




handle=bro.current_window_handle
html=bro.page_source 
soup=BeautifulSoup(html,'html.parser')
name_list = []
detail_list = []
price_list = []
count_list = []
web_list  = []
img_list  = []

# name
n = soup.find_all('a',class_="item-title J_GoldReport J_MakePoint")
ni = str(n)
nia = re.compile('>(.*?)</a>')
name_list=re.findall(nia,ni)
# print(*name_list)
#sys.stdout.flush()
# print(len(name_list))
#weblist
webs = soup.find_all('div',class_="item-info")
for web in webs:
    web_url = web.find('div',class_="item-basic-info")
    # print(web_url)
    # sys.stdout.flush()
    if web_url:
        website = web_url.a['href']
        web_list.append('https:'+ website+' ')
# for i in range(0,len(name_list)):
#     web_list.append(n[i]['href'])
# print(*web_list)
#sys.stdout.flush()
#details
# items = soup.find_all('div', class_="item-props item-props-can")
# for item in items:
#     color_element = item.find_all('p', class_="sku-line", tabindex="0")
#     if len(color_element) != 0 :
#          de = str( color_element)
#          det = re.compile('>(.*?)</p>')
#          detail_list.append( re.findall(det,de))
#     else:
#         detail_list.append(["无特殊备注"])

items = soup.find_all('li', class_="td td-info")
detail_list = []

for item in items:
    props = item.find('div', class_=re.compile(r"item-props(-can)?"))
    
    if props:
        color_elements = props.find_all('p', class_="sku-line", tabindex="0")
        if len(color_elements) > 0:
            details = [re.findall('>(.*?)</p>', str(color)) for color in color_elements]
            detail_list.append(details)
        else:
            detail_list.append("无特殊备注")
    else:
        detail_list.append("无特殊备注")

# d = soup.find_all('p',class_= "sku-line")
# de = str(d)
# det = re.compile('>(.*?)</p>')
# detail_list = re.findall(det,de)
# print(*detail_list)
#sys.stdout.flush()
# price
j=soup.find_all('em',class_="J_Price price-now")
js = str(j)
jia=re.compile('>(.*?)</em>')
price_list=re.findall(jia,js)
#print(*price_list)


# count 
counts = soup.find_all('div', class_="item-amount")
for count in counts:
    input_element = count.find('input', class_="text text-amount J_ItemAmount")
    if input_element:
        value = input_element['value']
        count_list.append(value)
    else:
        count_list.append("0")
###print(*count_list)
#sys.stdout.flush()

#image
images = soup.find_all('div',class_= "td-inner")
for image in images:
    img_url = image.find('img',class_= "itempic J_ItemImg")
    if img_url:
        img = img_url['src']
        img = img.replace("_80x80.jpg", "")
        img_list.append(img)
#print(*img_list)
#sys.stdout.flush()  

#print("\n\n\n\n",name_list[0],img_list[0], price_list[0],web_list[0],detail_list[0],count_list[0])
#sys.stdout.flush()
# r = soup.find_all('a',class_= 'item-basic-info')
# mi = str(r)
# sys.stdout.write(mi)
# #sys.stdout.flush()

#文档处理

# 获取桌面路径
desktop_path = os.path.expanduser("~\Desktop")

# 构造采购表的完整路径
docx_path = os.path.join(desktop_path, docx_name+'.docx')
# print(docx_path)
# sys.stdout.flush()
# 打开Word文档
doc = Document(docx_path)

# 获取第一个表格
table = doc.tables[0]
cursor = len(table.rows)
x_cor = len(table.rows)
y_cor = len(table.columns)
# 在表格末尾添加新行
for i in range(0,len(name_list)-1):
    text_progress_bar(len(name_list-1),i)
    sys.stdout.write("\n") 
    if float(count_list[i])!=0:

        #print(name_list[i],img_list[i], price_list[i],web_list[i],detail_list[i],count_list[i])
        #sys.stdout.flush()
        new_row = table.add_row().cells

        # 设置新行的单元格内容
        #print("开始添加")
        #sys.stdout.flush()
        detail_string =''
        for v in range(0,len(detail_list[i])):
            detail_string += str(detail_list[i][v])
            sys.stdout.write(detail_string + '\n')
            sys.stdout.flush()
        new_row[0].text = str(cursor)
        # new_row[1].text = '名称: '+str(name_list[i])+'\n型号: '+detail_string
        # paragraph = new_row[1].paragraphs[0]

        # # Modify the paragraph properties
        # paragraph_format = paragraph.paragraph_format
        # paragraph_format.space_after = 0

        # # Access the runs in the paragraph
        # for run in paragraph.runs:
        #     run.font.name = 'Arial'
        #     run.bold = True
        #     run.italic = True
        #     run.font.size = docx.shared.Pt(10)


        run = new_row[1].paragraphs[0].add_run('名称: '+str(name_list[i])+'\n型号: '+detail_string)
        run.font.name = 'Arial'
        run.bold = True
        run.italic = True
        run.font.size = docx.shared.Pt(10)
        
        # new_row[1].text = '名称: '+str(name_list[i]) +'\n型号: '+str(detail_list[i])
        # headers ={
        # 'User-Agent':'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/96.0.4664.45 Safari/537.36'
        # }
        # response = req.get(img_list[i],headers = headers)
        
        try:
            bro.get(web_list[i])
            time.sleep(3)
            path_verification2 = os.path.abspath(os.path.dirname(__file__))
            # print(path_verification)
            # sys.stdout.flush()
            slider_img2 = cv2.imread(path_verification2+"\dragslider2.png")
            slider_retry = cv2.imread(path_verification2+"\dragfailed.png")
            flag = False
            location_slider2 = pyautogui.locateOnScreen(slider_img2 ,confidence = 0.8)
            if location_slider2:
                
                point_slider2 = pyautogui.center(location_slider2)
                
                x_2, y_2 = point_slider2

                pyautogui.moveTo(x_2,y_2)
                pyautogui.mouseDown(x=x_2, y=y_2, button='left')
                rand_y = random.randint(0,5)
                rand_time = random.uniform(0, 1)
                pyautogui.moveTo(x=x_2+100, y=y_2+rand_y, duration=rand_time)
                
                rand_y = random.randint(0,5)
                rand_time = random.uniform(0, 1)
                pyautogui.moveTo(x=x_2+130, y=y_2+rand_y, duration=rand_time)
                rand_y = random.randint(5,10)
                rand_time = random.uniform(0, 1)
                pyautogui.moveTo(x=x_2+200, y=y_2+rand_y, duration=rand_time)
                time.sleep(random.uniform(0, 1))
                rand_y = random.randint(10,20)
                pyautogui.moveTo(x=x_2+235, y=y_2-rand_y, duration=random.uniform(0, 1))
                rand_y = random.randint(0,10)
                pyautogui.moveTo(x=x_2+350, y=y_2+rand_y, duration=random.uniform(0, 1))
                time.sleep(random.uniform(0, 1))
                pyautogui.mouseUp()
                time.sleep(0.5)
                flag = True
                # while_flag  = 1
                # try:
                location_failedslide = pyautogui.locateOnScreen(slider_retry, confidence = 0.8)
                if location_failedslide:
                    time.sleep(10)
                    flag = True
                #     while_flag = 1
                #     point_slider_failed = pyautogui.center(location_failedslide)
                #     x_3, y_3 = point_slider_failed
                #     pyautogui.moveTo(x_3, y_3)
                #     pyautogui.click()
                #     pyautogui.moveTo(x_2,y_2)
                #     pyautogui.mouseDown(x=x_2, y=y_2, button='left')
                #     pyautogui.moveTo(x=x_2+100, y=y_2, duration=0.2)
                #     time.sleep(0.2)
                #     pyautogui.moveTo(x=x_2+200, y=y_2, duration=0.2)
                #     time.sleep(0.2)
                #     pyautogui.moveTo(x=x_2+230, y=y_2-10, duration=1)
                #     time.sleep(0.2)
                #     pyautogui.moveTo(x=x_2+350, y=y_2+5, duration=0.4)
                #     pyautogui.mouseUp()

                # except:
                #     # break
                    
                #     sys.stdout.write("没有重试\n")
            else:
                flag = True
                sys.stdout.write("没有滑动条\n")
            while not flag:
                pass
            flag = False
            #模拟预选商品信息
            bro.implicitly_wait(5)
            html_content = bro.page_source

            # 将HTML内容转换为字符串
            html_string = str(html_content)

            # 要搜索的目标字符串
            # element_locator1 = (By.ID, "skuItem")
            # element_locator2 = (By.ID, "//s.taobao.com/search")
            # wait = WebDriverWait(bro, 10)
            # wait.until(EC.presence_of_element_located(
            #     tuple(chain(element_locator1, element_locator2)))
            # )

            target_string = "skuItem"
            target_string2 = "//s.taobao.com/search"

            # 检查目标字符串是否在HTML中存在
            if target_string in html_string:
                # 在HTML中找到目标字符串，模拟鼠标点击该元素

                print("天猫")
                titlelist = []
                sys.stdout.flush()
                soup1=BeautifulSoup(html_content,'html.parser')
                # sku_items = soup1.find_all('div',class_= 'skuValueName')
                #

                # Process each skuItem element
                # v = soup.find_all('div',class_='skuItem')
                # vi = str(v)
                # # via = re.compile('>(.*?)</div>')
                # # titlelist=re.findall(via,vi)
                # print(vi)
                # sys.stdout.flush()

                def filter_sku_item(tag):
                    return tag.name == 'span' and tag.get('class') == ['skuValueName']
                sku_items = soup1.find_all(filter_sku_item)
                # l=bro.find_element(By.XPATH,"//div[@title='爆款 新品升级【均衡之刃】1支单拍 训练专用 亚光红5u 球队训练比赛']")
                # l.click()
                for sku_item in sku_items:
                    for j in range(len(detail_list[i])):
                        detail_list_remove = str(detail_list[i][j])
                        detail_list_remove = detail_list_remove.replace('[','')
                        detail_list_remove = detail_list_remove.replace(']','')
                        detail_list_remove = detail_list_remove.replace(" ","")
                        sku_item_list  =  str(*sku_item).replace(" ","")
                        sku_item_string = str(*sku_item)
                        if find_lc(detail_list_remove,sku_item_list) == len(sku_item_list):
                            l = bro.find_element(By.XPATH, '//div[@title="{}"]'.format(sku_item_string))
                            l.click()
                            time.sleep(1)
                            
                        # print(find_lc(detail_list_remove,sku_item_list))
                        # print(len(sku_item_list))
                        # print(detail_list_remove)
                        # print(sku_item_list)
                        # sys.stdout.flush()
                url_tianmao = bro.current_url
                web_list[i] = str(url_tianmao)
                print(str(url_tianmao))
                sys.stdout.flush()
                flag = True
                # time.sleep(1)
                # bro.quit()
                # sys.exit(0)
                # for sku_item in sku_items:
                # # Get the string after the 'title' attribute within the 'div' element
                #     title_string = sku_item['title']
                #     text_after_title = sku_item.text.replace(title_string, '', 1).strip()
                #     print(text_after_title)
                #     sys.stdout.flush()
                # sample_sku_item = sku_items[0]
                # title_element = sample_sku_item.get('title')
                # print(title_element)
                # sys.stdout.flush()
                # for sku_item in sku_items:
                #     title_element = sku_item.get('title')
                #     print(title_element)
                #     sys.stdout.flush()
                    # if title_element: 
                    #     if has_common_elements(detail_list[i],title_element):
                    #         # 在detail_list中找到匹配的内容，模拟鼠标点击该<div>元素
                    #         # Your code to simulate the click here
                    #         title_element.click()
                    #         sys.stdout.write("预选完成")
                    #         sys.stdout.flush()
            elif target_string2 in html_string:
                print("淘宝")
                # try:
                title_list=[]
                soup2=BeautifulSoup(html_content,'html.parser')
                titles = soup2.find_all('ul',class_= "J_TSaleProp tb-img tb-clearfix")
                prefixes = soup2.find_all('dt', class_='tb-property-type')
                #[<dt class="tb-property-type">颜色分类</dt>, <dt class="tb-property-type">数量</dt>]
                prefix_texts = [prefix.text for prefix in prefixes]
                # 输出结果
                # print(prefix_texts)
                titles_str = str(titles)
                for item in prefix_texts:
                    titles_str = titles_str.replace(item, "")
                    print("最长相同部分为:", titles_str) 
                    sys.stdout.flush()
                for j in range(len(detail_list[i])):
                    print(j)
                    print(str(detail_list[i][j]))
                    sys.stdout.flush()
                    result = longest_common_substring(str(detail_list[i][j]),titles_str,"颜色分类")
                    print("最长相同部分为:", result) 
                    sys.stdout.flush()
                    if result in titles_str:
                        try:
                            v=bro.find_element(By.XPATH, "//li[./a/span[text()='{}']]".format(result))
                            v.click()
                            # if j == range(len(detail_list[i])):
                            #     flag = True
                        except:
                            print('没找到')
                            sys.stdout.flush()
                flag = True
                # except:
                #     flag = True
                #     sys.stdout.write("淘宝分类跳出 直接截图\n")
                # bro.quit()
                # sys.exit(0)
            else:
                flag = True
                print("目标字符串在HTML中未找到。")
            # sku_items = soup.find_all('div',class_ = "skuItem ")
            # for sku_item in sku_items:
            #     title_element = sku_item['title']
            #     if title_element: 
            #         if has_common_elements(detail_list[i],title_element):
            #             # 在detail_list中找到匹配的内容，模拟鼠标点击该<div>元素
            #             # Your code to simulate the click here
            #             title_element.click()
            #             sys.stdout.write("预选完成")
            #             sys.stdout.flush()
            #     else:
            #         sys.stdout.write("预选失败\n")
            #         sys.stdout.flush()     
            #预选完成
            while not flag:
                time.sleep(0.1)
            print("截图中")
            sys.stdout.flush()
            window_width = bro.execute_script("return window.innerWidth")
            window_height = bro.execute_script("return window.innerHeight")

            # 计算截图区域的左上角和右下角坐标
            left = int(window_width * 0.125)  # 屏幕宽度的12.5%
            top = int(window_height * 0.125)  # 屏幕高度的12.5%
            right = int(window_width * 0.85)  # 屏幕宽度的85%
            bottom = int(window_height)  # 屏幕高度的100%      
            screenshot = bro.get_screenshot_as_png()
            # 使用imageio库加载图像

            image = Image.open(io.BytesIO(screenshot))
            cropped_image = image.crop((left, top, right, bottom))
            image_folder = os.path.expanduser(r"~\Desktop\taobao")
            image_path = os.path.join(image_folder,f"image_{i}.png")
            cropped_image.show()
            cropped_image.save(image_path)
            # 在表格中添加一行
            #print(" error")
            new_row[2].paragraphs[0].add_run().add_picture(image_path, width=Inches(1.0))
            os.remove(image_path)
            print("截图完成")
            sys.stdout.flush() 
            #print("No error")

        except Exception as e:
            print(f"Error loading image")
            print(e)
        p = new_row[3].add_paragraph()
        # new_row[3].text = web_list[i]
        add_hyperlink(p,web_list[i],web_list[i])
        price_int  =int(float(price_list[i][1:])/float(count_list[i])) 
        new_row[4].text = str(price_int)+'元/件'
        new_row[5].text = count_list[i]
        cursor += 1
        # text_progress_bar(len(name_list),i)
        #print("结束添加")
        #sys.stdout.flush()
# 保存修改后的文档
modified_docx_path = os.path.join(desktop_path, docx_name+".docx")
doc.save(modified_docx_path)
bro.quit()
