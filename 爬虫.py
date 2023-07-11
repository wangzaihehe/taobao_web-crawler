import pyautogui
import sys
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from bs4 import BeautifulSoup
import time
import cv2
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.support.ui import  WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import re

from docx import Document
import docx
import os
import requests as req
from PIL import Image
import io
from io import BytesIO
from docx.shared import Inches
import imageio
import magic
import numpy as np
def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = docx.text.run.Run(
        docx.oxml.shared.OxmlElement('w:r'), paragraph)
    new_run.text = text

    # Set the run's style to the builtin hyperlink style, defining it if necessary
    # new_run.style = get_or_create_hyperlink_style(part.document)
    # Alternatively, set the run's formatting explicitly
    new_run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
    new_run.font.underline = True

    # Join all the xml elements together
    hyperlink.append(new_run._element)  
    paragraph._p.append(hyperlink)
    return hyperlink


# chrome_img = cv2.imread(r"C:\Users\WIN10\Desktop\chrome1.png") 
# pyautogui.hotkey('win','m')
# location_chrome = pyautogui.locateOnScreen(chrome_img,confidence = 0.7)
# print(location_chrome)
# pyautogui.doubleClick(location_chrome)
# searchbar_img = cv2.imread(r"C:\Users\WIN10\Desktop\searchbar.png") 
# location_searchbar = pyautogui.locateOnScreen(searchbar_img ,confidence = 0.7)
# print(location_searchbar)
# pyautogui.click(location_searchbar)
# pyautogui.write('https://login.taobao.com/')
docx_name = pyautogui.prompt(text='请输入位于桌面的文档名「无需后缀」',title='购物车爬虫消息框',default='')
id = pyautogui.prompt(text='请输入淘宝用户名',title='购物车爬虫消息框',default='')
password = pyautogui.password(text='输入密码',title='购物车爬虫消息框',default='',mask='*')
pyautogui.alert(text='在登陆后如遇二重验证,请在15秒内完成手机确认',title='购物车爬虫消息框',button='OK')
# id ='大穷逼汪仔'
# password = 'wangzai123'
option = Options()
option.add_experimental_option('excludeSwitches', ['enable-automation'])
option.add_argument('--disable-blink-features=AutomationControlled')
bro=webdriver.Chrome(options = option)
bro.maximize_window()
bro.implicitly_wait(5)
wait = WebDriverWait(bro, 10)
bro.get('https://login.taobao.com/')
# login_img = cv2.imread(r"C:\Users\WIN10\Desktop\username.png") 
# location_login = pyautogui.locateOnScreen(login_img)
# print(location_login)

# pyautogui.typewrite('Hello world')
# driver = webdriver.Chrome()

bro.find_element(By.ID,'fm-login-id').click()
bro.find_element(By.ID,'fm-login-id').clear()
bro.find_element(By.ID,'fm-login-id').send_keys(id)

bro.find_element(By.ID,'fm-login-password').clear()
bro.find_element(By.ID,'fm-login-password').send_keys(password)
bro.find_element(By.ID,'fm-login-password').send_keys(Keys.ENTER)
time.sleep(4)
# bro.find_element(By.CLASS_NAME,'fm-button.fm-submit.password-login').click()
#滑块验证
path_verification = os.path.abspath(os.path.dirname(__file__))
# print(path_verification)
# sys.stdout.flush()
slider_img = cv2.imread(path_verification+"\dragslider.png") 
try:

    location_slider = pyautogui.locateOnScreen(slider_img ,confidence = 0.8)
    
    point_slider = pyautogui.center(location_slider)
    time.sleep(4)
    x, y = point_slider

    pyautogui.moveTo(x,y)
    pyautogui.dragTo(x+300, y, 0.5, button='left') 
except:
    # sys.stdout.write("没有滑动条")
    time.sleep(1)

time.sleep(3)
handle=bro.current_window_handle
bro.find_element(By.ID,'J_MiniCart').click()
bro.implicitly_wait(5)
wait = WebDriverWait(bro, 10)
#滚动
scroll_pause_time = 5 
last_height = bro.execute_script('return document.body.scrollHeight')
while True:
    bro.find_element('tag name', 'body').send_keys(Keys.END)
    time.sleep(scroll_pause_time)
    new_height = bro.execute_script('return document.body.scrollHeight')
    if new_height == last_height:
        break
    last_height = new_height
bro.implicitly_wait(5)
#全部界面加载完成
handle=bro.current_window_handle
html=bro.page_source 
soup=BeautifulSoup(html,'html.parser')
name_list = []
detail_list = []
price_list = []
count_list = []
web_list  = []
img_list  = []

# name
n = soup.find_all('a',class_="item-title J_GoldReport J_MakePoint")
ni = str(n)
nia = re.compile('>(.*?)</a>')
name_list=re.findall(nia,ni)
# print(*name_list)
#sys.stdout.flush()
# print(len(name_list))
#weblist
webs = soup.find_all('div',class_="item-info")
for web in webs:
    web_url = web.find('div',class_="item-basic-info")
    if web_url:
        website = web_url.a['href']
        web_list.append('https:'+ website+' ')
# for i in range(0,len(name_list)):
#     web_list.append(n[i]['href'])
# print(*web_list)
#sys.stdout.flush()
#details
# items = soup.find_all('div', class_="item-props item-props-can")
# for item in items:
#     color_element = item.find_all('p', class_="sku-line", tabindex="0")
#     if len(color_element) != 0 :
#          de = str( color_element)
#          det = re.compile('>(.*?)</p>')
#          detail_list.append( re.findall(det,de))
#     else:
#         detail_list.append(["无特殊备注"])

items = soup.find_all('li', class_="td td-info")
detail_list = []

for item in items:
    props = item.find('div', class_=re.compile(r"item-props(-can)?"))
    
    if props:
        color_elements = props.find_all('p', class_="sku-line", tabindex="0")
        if len(color_elements) > 0:
            details = [re.findall('>(.*?)</p>', str(color)) for color in color_elements]
            detail_list.append(details)
        else:
            detail_list.append("无特殊备注")
    else:
        detail_list.append("无特殊备注")

# d = soup.find_all('p',class_= "sku-line")
# de = str(d)
# det = re.compile('>(.*?)</p>')
# detail_list = re.findall(det,de)
# print(*detail_list)
#sys.stdout.flush()
# price
j=soup.find_all('em',class_="J_Price price-now")
js = str(j)
jia=re.compile('>(.*?)</em>')
price_list=re.findall(jia,js)
#print(*price_list)


# count 
counts = soup.find_all('div', class_="item-amount")
for count in counts:
    input_element = count.find('input', class_="text text-amount J_ItemAmount")
    if input_element:
        value = input_element['value']
        count_list.append(value)
    else:
        count_list.append("0")
###print(*count_list)
#sys.stdout.flush()

#image
images = soup.find_all('div',class_= "td-inner")
for image in images:
    img_url = image.find('img',class_= "itempic J_ItemImg")
    if img_url:
        img = img_url['src']
        img = img.replace("_80x80.jpg", "")
        img_list.append(img)
#print(*img_list)
#sys.stdout.flush()  

#print("\n\n\n\n",name_list[0],img_list[0], price_list[0],web_list[0],detail_list[0],count_list[0])
#sys.stdout.flush()
# r = soup.find_all('a',class_= 'item-basic-info')
# mi = str(r)
# sys.stdout.write(mi)
# #sys.stdout.flush()

#文档处理

# 获取桌面路径
desktop_path = os.path.expanduser("~/Desktop")

# 构造采购表的完整路径
docx_path = os.path.join(desktop_path, docx_name+'.docx')

# 打开Word文档
doc = Document(docx_path)

# 获取第一个表格
table = doc.tables[0]
cursor = len(table.rows)
x_cor = len(table.rows)
y_cor = len(table.columns)
# 在表格末尾添加新行
for i in range(0,len(name_list)-1):
    if float(count_list[i])!=0:

        #print(name_list[i],img_list[i], price_list[i],web_list[i],detail_list[i],count_list[i])
        #sys.stdout.flush()
        new_row = table.add_row().cells

        # 设置新行的单元格内容
        #print("开始添加")
        #sys.stdout.flush()
        detail_string =''
        for v in range(0,len(detail_list[i])):
            detail_string += str(detail_list[i][v])
            detail_string += ' '
        new_row[0].text = str(cursor)
        run = new_row[1].add_paragraph().add_run('名称: '+str(name_list[i])+'\n型号: '+detail_string)
        run.font.name = 'Arial'
        run.bold = True
        run.italic = True
        run.font.size = docx.shared.Pt(10)
        
        # new_row[1].text = '名称: '+str(name_list[i]) +'\n型号: '+str(detail_list[i])
        headers ={
        'User-Agent':'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/96.0.4664.45 Safari/537.36'
        }
        response = req.get(img_list[i],headers = headers)
        
        try:

            # 使用imageio库加载图像

            image = Image.open(BytesIO(response.content))
            if image.mode in ("RGBA", "P"): image = image.convert("RGB")
            image_path = f"image_{i}.jpg"
            image.save(image_path, "JPEG")
            # 在表格中添加一行
            #print(" error")
            new_row[2].paragraphs[0].add_run().add_picture(image_path, width=Inches(1.0))
            os.remove(image_path) 
            #print("No error")

        except Exception as e:
            print(f"Error loading image: {img_list[i]}")
            print(e)
        p = new_row[3].add_paragraph()
        # new_row[3].text = web_list[i]
        add_hyperlink(p,web_list[i],web_list[i])
        price_int  = float(price_list[i][1:])/float(count_list[i])
        new_row[4].text = str(price_int)+'元/件'
        new_row[5].text = count_list[i]
        cursor += 1
        
        #print("结束添加")
        #sys.stdout.flush()
# 保存修改后的文档
modified_docx_path = os.path.join(desktop_path, docx_name+".docx")
doc.save(modified_docx_path)
bro.quit()
